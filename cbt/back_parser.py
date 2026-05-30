"""
Robust Word Document Parser for CBT Questions.

Accepts ALL common teacher formats — no training needed.

QUESTION LINE — any of:
    1. text        1) text       Q1. text      Q1) text
    Q 1. text      Question 1.   (1) text      1- text
    1: text        40. text      Q1 text (no delimiter)

OPTION LINE — any of:
    A. text   A) text   A: text   A- text
    a. text   (A) text  A.text (no space)

ANSWER LINE — any of:
    ANSWER: A    Answer A     Ans: B      ANS B
    Answer:C     ANSWER:A     answer = A  ANSWER- A
    Correct Answer: D         The answer is A   Key: B

MARKS LINE — any of:
    MARKS: 1    MARKS: 0.5    MARKS: 1/2    MARKS: 2/4
    POINTS: 2   MARK: .5      score: 1.5    marks 2

INLINE (entire question on one paragraph):
    40. According to ionic theory, acids produce A. OH⁻ ions B. H⁺ ions C. Na⁺ ions D. Cl⁻ ions ANSWER: B MARK: 0.5
    Q1. Capital city?A.LagosB.AbujaC.KanoD.IbadanANSWER:B
"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from io import BytesIO
import re
import hashlib


# ══════════════════════════════════════════════════════════════
#  COMPILED PATTERNS
# ══════════════════════════════════════════════════════════════

# Question: number with any delimiter (or none), with optional Q/Question prefix
RE_QUESTION = re.compile(
    r'^(?:\(?\s*Q(?:uestion)?\s*\.?\s*)?'   # optional Q / Question prefix
    r'\(?\s*(\d+)\s*[.):\-]?\s*\)?\s*'      # the number + any delimiter
    r'(.+)',                                   # the question text
    re.IGNORECASE | re.DOTALL,
)

# Option line: A. / A) / A: / A- / (A) / a. — all accepted
RE_OPTION_LINE = re.compile(
    r'^\(?([A-D])[.):\-]\)?\s*(.*)',
    re.IGNORECASE | re.DOTALL,
)

# Answer at the START of a line
RE_ANSWER = re.compile(
    r'^(?:the\s+)?'
    r'(?:ANSWER|ANS(?:WER)?|CORRECT\s*ANSWER|KEY)'
    r'(?:\s+is)?'
    r'[\s:=\-]*([A-D])\b',
    re.IGNORECASE,
)

# Answer ANYWHERE inside a string (for inline scanning)
RE_ANSWER_SEARCH = re.compile(
    r'(?:ANSWER|ANS(?:WER)?|CORRECT\s*ANSWER|KEY)'
    r'(?:\s+is)?'
    r'[\s:=\-]*([A-D])\b',
    re.IGNORECASE,
)

# Marks at the START of a line
RE_MARKS = re.compile(
    r'^(?:MARKS?|POINTS?|SCORE|MARK|WEIGHT)[:\s=]+(.+)',
    re.IGNORECASE,
)

# Marks ANYWHERE inside a string (for inline scanning)
RE_MARKS_SEARCH = re.compile(
    r'(?:MARKS?|POINTS?|SCORE|MARK|WEIGHT)[:\s=]+([0-9/.]+)',
    re.IGNORECASE,
)

# Splits inline options at every A. / A) boundary
RE_OPT_SPLIT = re.compile(r'(?=[A-D][.)])', re.IGNORECASE)

# Detects inline options (spaced: " A. text")
RE_INLINE_OPT_SPACED = re.compile(r'(?:^|\s)([A-D])[.)]\s*\S', re.IGNORECASE)

# Detects inline options (glued after punctuation: "textA. text")
RE_INLINE_OPT_GLUED  = re.compile(r'(?<=[a-z\d?!,;])([A-D])[.)]', re.IGNORECASE)


# ══════════════════════════════════════════════════════════════
#  HELPERS
# ══════════════════════════════════════════════════════════════

def _parse_marks(raw):
    """
    Parse marks from raw string.
    Handles: "1", "0.5", ".5", "1/2", "3/4", "2/4", "1.5"
    Returns float, defaulting to 1.0 on failure.
    """
    if not raw:
        return 1.0
    raw = raw.strip()
    if '/' in raw:
        parts = raw.split('/', 1)
        try:
            return round(float(parts[0].strip()) / float(parts[1].strip()), 4)
        except (ValueError, ZeroDivisionError):
            pass
    try:
        return float(raw)
    except ValueError:
        return 1.0


def _find_inline_options_start(text):
    """
    Return the index where inline options begin inside `text`,
    or -1 if no inline options are found.
    Requires at least 2 distinct option letters to avoid false positives.
    """
    candidates = []
    for m in RE_INLINE_OPT_SPACED.finditer(text):
        candidates.append((m.start(), m.group(1).upper()))
    for m in RE_INLINE_OPT_GLUED.finditer(text):
        candidates.append((m.start(), m.group(1).upper()))

    if not candidates:
        return -1
    distinct_letters = {c[1] for c in candidates}
    if len(distinct_letters) < 2:
        return -1
    return min(c[0] for c in candidates)


def _is_inline_question(text):
    """
    True when a single paragraph contains a question number AND
    embedded options (A. / A) appearing at least twice).
    """
    if not RE_QUESTION.match(text):
        return False
    q_m = RE_QUESTION.match(text)
    rest = text[q_m.start(2):]
    return _find_inline_options_start(rest) >= 0


def _extract_inline_options(options_text):
    """
    Parse the options+answer+marks section of an inline question.

    e.g. " A. OH⁻ ions B. H⁺ ions C. Na⁺ ions D. Cl⁻ ions ANSWER: B MARK: 0.5"

    Returns: (options_dict, answer_letter_or_None, marks_float_or_None)
    """
    options = {}
    answer  = None
    marks   = None

    for chunk in RE_OPT_SPLIT.split(options_text):
        chunk = chunk.strip()
        if not chunk:
            continue

        m = re.match(r'^([A-D])[.)]\s*(.*)', chunk, re.IGNORECASE | re.DOTALL)
        if not m:
            # Stray text — may contain ANSWER or MARKS
            a = RE_ANSWER_SEARCH.search(chunk)
            if a:
                answer = a.group(1).upper()
            mk = RE_MARKS_SEARCH.search(chunk)
            if mk:
                marks = _parse_marks(mk.group(1))
            continue

        letter      = m.group(1).upper()
        option_text = m.group(2).strip()

        # ── Scan for MARKS first (before we truncate at ANSWER) ──────────
        mk = RE_MARKS_SEARCH.search(option_text)
        if mk:
            marks = _parse_marks(mk.group(1))

        # ── Scan for ANSWER ───────────────────────────────────────────────
        a = RE_ANSWER_SEARCH.search(option_text)
        if a:
            answer      = a.group(1).upper()
            # Truncate option text at whichever comes first: ANSWER or MARKS
            cut = a.start()
            if mk and mk.start() < cut:
                cut = mk.start()
            option_text = option_text[:cut].strip()
        elif mk:
            option_text = option_text[:mk.start()].strip()

        if option_text:
            options[letter] = option_text

    return options, answer, marks


def _extract_paragraph_image(paragraph, document):
    """
    If a paragraph contains an embedded image (pasted directly into Word),
    extract the raw image bytes and a generated filename.
    Returns (filename, bytes) or (None, None).
    """
    try:
        # Look for <w:drawing> elements in the paragraph XML
        drawing_els = paragraph._element.findall('.//' + qn('w:drawing'))
        if not drawing_els:
            return None, None

        # Find the relationship ID of the image blip
        blip = None
        for drawing in drawing_els:
            blip = drawing.find('.//' + qn('a:blip'))
            if blip is not None:
                break

        if blip is None:
            return None, None

        # Get the rId (relationship id)
        r_id = blip.get(qn('r:embed'))
        if not r_id:
            return None, None

        # Get image bytes via the document's part relationships
        image_part = paragraph.part.related_parts.get(r_id)
        if image_part is None:
            return None, None

        image_bytes = image_part.blob
        # Generate a stable filename from content hash + extension
        ext = image_part.content_type.split('/')[-1]
        if ext == 'jpeg':
            ext = 'jpg'
        fname = f"embed_{hashlib.md5(image_bytes[:64]).hexdigest()[:10]}.{ext}"
        return fname, image_bytes

    except Exception:
        return None, None


# ══════════════════════════════════════════════════════════════
#  MAIN PARSER CLASS
# ══════════════════════════════════════════════════════════════

class WordQuestionParser:
    """
    Robust parser for Word documents containing CBT questions in any format.
    """

    def __init__(self, file_path_or_stream):
        self.document  = Document(file_path_or_stream)
        self.questions = []
        self.current_passage       = None
        self.current_passage_title = None
        self.current_passage_group = None

    # ── Public ────────────────────────────────────────────────────────────

    def parse(self):
        current_question = None
        current_options  = {}

        for paragraph in self.document.paragraphs:
            raw_text = paragraph.text.strip()

            # ── Embedded image (pasted directly into Word) ────────────────
            # Check BEFORE text processing — images may sit in otherwise-empty
            # paragraphs or alongside question text
            emb_fname, emb_bytes = _extract_paragraph_image(paragraph, self.document)
            if emb_fname and emb_bytes:
                if current_question and not current_question.get('image_reference'):
                    current_question['image_reference'] = emb_fname
                    current_question['image_bytes']     = emb_bytes

            # Strip any IMAGE:/CAPTION: markers that crept into question text
            # (happens when IMAGE: appears on the same paragraph as the question)
            text = re.sub(r'\s*IMAGE\s*:\s*\S+', '', raw_text, flags=re.IGNORECASE).strip()
            text = re.sub(r'\s*CAPTION\s*:\s*.+$', '', text, flags=re.IGNORECASE).strip()

            if not text:
                continue

            # ── Passage markers ───────────────────────────────────────────
            pg = re.match(r'^PASSAGE[_\s]?GROUP[:\s]+(.+)', text, re.IGNORECASE)
            if pg:
                self.current_passage_group = pg.group(1).strip()
                continue

            pt = re.match(r'^PASSAGE[:\s]+(.+)', text, re.IGNORECASE)
            if pt:
                self.current_passage_title = pt.group(1).strip()
                continue

            if re.match(r'^START[_\s]?PASSAGE', text, re.IGNORECASE):
                self.current_passage = []
                continue

            if re.match(r'^END[_\s]?PASSAGE', text, re.IGNORECASE):
                if isinstance(self.current_passage, list):
                    self.current_passage = '\n'.join(self.current_passage)
                continue

            if isinstance(self.current_passage, list):
                self.current_passage.append(text)
                continue

            # ── Image / caption markers ───────────────────────────────────
            img = re.match(r'^IMAGE[:\s]+(.+)', text, re.IGNORECASE)
            if img:
                ref = img.group(1).strip()
                if current_question:
                    current_question['image_reference'] = ref
                else:
                    # IMAGE: appeared before the question number line — store
                    # it on _pending so _new_question can pick it up
                    self._pending_image_ref = ref
                continue

            cap = re.match(r'^CAPTION[:\s]+(.+)', text, re.IGNORECASE)
            if cap:
                if current_question:
                    current_question['image_caption'] = cap.group(1).strip()
                else:
                    self._pending_image_cap = cap.group(1).strip()
                continue

            # ── Marks line ────────────────────────────────────────────────
            mk = RE_MARKS.match(text)
            if mk and current_question:
                current_question['marks'] = _parse_marks(mk.group(1))
                continue

            # ── Short-answer marker ───────────────────────────────────────
            if re.match(r'^(?:SHORT\s*ANSWER|SA)[:\s]*$', text, re.IGNORECASE) and current_question:
                current_question['question_type'] = 'SA'
                continue

            sa = re.match(r'^(?:ANSWER|ANS)[:\s]+(.+)', text, re.IGNORECASE)
            if sa and current_question and current_question.get('question_type') == 'SA':
                current_question['correct_answer'] = sa.group(1).strip()
                continue

            # ══ INLINE QUESTION ══════════════════════════════════════════
            # Everything (question text + A/B/C/D options + answer + marks)
            # is on a SINGLE paragraph line.
            if _is_inline_question(text):
                # Save any in-progress question first
                if current_question:
                    current_question['options'] = current_options
                    self.questions.append(current_question)
                    current_question = None
                    current_options  = {}

                q_m       = RE_QUESTION.match(text)
                full_rest = text[q_m.start(2):]
                opt_idx   = _find_inline_options_start(full_rest)

                # Split: question text | options+answer+marks
                question_text = full_rest[:opt_idx].strip() if opt_idx > 0 else full_rest.strip()
                options_text  = full_rest[opt_idx:]         if opt_idx > 0 else ''

                opts, ans, inline_marks = _extract_inline_options(options_text)

                q = self._new_question(q_m.group(1), question_text)
                q['options'] = opts
                if ans:
                    q['correct_option'] = ans
                if inline_marks is not None:
                    q['marks'] = inline_marks

                self.questions.append(q)
                continue   # inline question fully handled — move on

            # ══ NORMAL QUESTION LINE ════════════════════════════════════
            q_m = RE_QUESTION.match(text)
            if q_m:
                # Save previous question
                if current_question:
                    current_question['options'] = current_options
                    self.questions.append(current_question)

                # Build HTML only for the question text portion
                q_html = self._paragraph_to_html(paragraph, strip_q_prefix=True)
                current_question = self._new_question(q_m.group(1), q_html)
                current_options  = {}
                continue

            # ── Option line ───────────────────────────────────────────────
            opt_m = RE_OPTION_LINE.match(text)
            if opt_m and current_question:
                letter   = opt_m.group(1).upper()
                opt_html = self._paragraph_to_html(paragraph, strip_opt_prefix=True)
                current_options[letter] = opt_html
                continue

            # ── Answer line ───────────────────────────────────────────────
            ans_m = RE_ANSWER.match(text)
            if ans_m and current_question:
                current_question['correct_option'] = ans_m.group(1).upper()
                # Auto-detect True/False
                if len(current_options) == 2:
                    vals = [str(v).lower() for v in current_options.values()]
                    if 'true' in vals and 'false' in vals:
                        current_question['question_type'] = 'TF'
                continue

        # Save last pending question
        if current_question:
            current_question['options'] = current_options
            self.questions.append(current_question)

        return self.questions

    # ── Helpers ───────────────────────────────────────────────────────────

    def _new_question(self, number, text):
        q = {
            'question_number': number,
            'question_text':   text,
            'question_type':   'MCQ',
            'marks':           1.0,
            'correct_option':  None,
            'correct_answer':  None,
            'passage':         self.current_passage if isinstance(self.current_passage, str) else None,
            'passage_title':   self.current_passage_title,
            'passage_group':   self.current_passage_group,
            'image_reference': getattr(self, '_pending_image_ref', None),
            'image_caption':   getattr(self, '_pending_image_cap', None),
            'image_bytes':     None,
        }
        self._pending_image_ref = None
        self._pending_image_cap = None
        return q

    def _paragraph_to_html(self, paragraph, strip_q_prefix=False, strip_opt_prefix=False):
        """Convert paragraph to HTML preserving bold/italic/super/subscript."""
        html_parts = []

        for run in paragraph.runs:
            text = run.text
            if not html_parts:
                if strip_q_prefix:
                    text = re.sub(
                        r'^(?:\(?\s*Q(?:uestion)?\s*\.?\s*)?\(?\s*\d+\s*[.):\-]?\s*\)?\s*',
                        '', text, flags=re.IGNORECASE
                    )
                elif strip_opt_prefix:
                    text = re.sub(
                        r'^\(?[A-D][.):\-]\)?\s*',
                        '', text, flags=re.IGNORECASE
                    )
            if not text:
                continue

            html = text
            if run.font.subscript:   html = f'<sub>{html}</sub>'
            if run.font.superscript: html = f'<sup>{html}</sup>'
            if run.bold:             html = f'<strong>{html}</strong>'
            if run.italic:           html = f'<em>{html}</em>'
            if run.underline:        html = f'<u>{html}</u>'
            html_parts.append(html)

        return ''.join(html_parts)

    def validate_questions(self):
        errors = []
        for idx, q in enumerate(self.questions, 1):
            if not q.get('question_text'):
                errors.append(f"Question {idx}: Missing question text")
            if q['question_type'] in ['MCQ', 'TF']:
                if not q.get('options') or len(q['options']) < 2:
                    errors.append(f"Question {idx}: Needs at least 2 options")
                if not q.get('correct_option'):
                    errors.append(f"Question {idx}: Missing correct answer")
                if q.get('correct_option') and q['correct_option'] not in q.get('options', {}):
                    errors.append(
                        f"Question {idx}: Answer '{q['correct_option']}' "
                        f"not in options {list(q.get('options', {}).keys())}"
                    )
            elif q['question_type'] == 'SA':
                if not q.get('correct_answer'):
                    errors.append(f"Question {idx}: Missing answer")
        return errors

    def get_statistics(self):
        stats = {
            'total': len(self.questions), 'mcq': 0, 'tf': 0, 'sa': 0,
            'total_marks': 0.0, 'with_passage': 0, 'with_image': 0,
            'passage_groups': set(),
        }
        for q in self.questions:
            t = q['question_type']
            if t == 'MCQ':  stats['mcq'] += 1
            elif t == 'TF': stats['tf']  += 1
            elif t == 'SA': stats['sa']  += 1
            stats['total_marks'] += q.get('marks', 1.0)
            if q.get('passage'):         stats['with_passage'] += 1
            if q.get('image_reference'): stats['with_image']   += 1
            if q.get('passage_group'):   stats['passage_groups'].add(q['passage_group'])
        stats['passage_groups'] = len(stats['passage_groups'])
        return stats


# ══════════════════════════════════════════════════════════════
#  TEMPLATE GENERATOR
# ══════════════════════════════════════════════════════════════

def create_word_template():
    doc = Document()

    t = doc.add_paragraph('CBT QUESTIONS TEMPLATE')
    t.runs[0].bold = True
    t.runs[0].font.size = Pt(16)

    doc.add_paragraph('All formats below are accepted:')
    for line in [
        'Questions : 1. text  |  1) text  |  Q1. text  |  (1) text  |  Q1 text',
        'Options   : A. text  |  A) text  |  (A) text  |  a. text   |  A.text',
        'Answers   : ANSWER: A  |  Answer A  |  Ans: B  |  The answer is C',
        'Marks     : MARKS: 1   |  MARKS: 0.5  |  MARKS: 1/2  |  POINTS: 2',
        'Inline    : 40. Question A. Opt1 B. Opt2 C. Opt3 D. Opt4 ANSWER: A MARKS: 0.5',
    ]:
        doc.add_paragraph(line)

    doc.add_paragraph()
    doc.add_paragraph('=== Normal format ===').runs[0].bold = True
    doc.add_paragraph('1. The capital of Nigeria is')
    doc.add_paragraph('A. Lagos')
    doc.add_paragraph('B. Abuja')
    doc.add_paragraph('C. Kano')
    doc.add_paragraph('D. Ibadan')
    doc.add_paragraph('ANSWER: B')
    doc.add_paragraph('MARKS: 0.5')

    doc.add_paragraph()
    doc.add_paragraph('=== Inline format ===').runs[0].bold = True
    doc.add_paragraph('40. According to ionic theory, acids produce A. OH⁻ ions B. H⁺ ions C. Na⁺ ions D. Cl⁻ ions ANSWER: B MARK: 0.5')

    doc.add_paragraph()
    doc.add_paragraph('=== Fraction marks ===').runs[0].bold = True
    doc.add_paragraph('2. What is H2O?')
    doc.add_paragraph('A. Salt')
    doc.add_paragraph('B. Water')
    doc.add_paragraph('C. Sugar')
    doc.add_paragraph('D. Oil')
    doc.add_paragraph('ANSWER: B')
    doc.add_paragraph('MARKS: 1/2')

    return doc