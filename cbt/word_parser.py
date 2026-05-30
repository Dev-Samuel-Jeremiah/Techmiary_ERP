"""
Robust Word Document Parser for CBT Questions.

Accepts ALL common teacher formats — no training needed.

IMAGE SUPPORT:
    1. Embedded/pasted directly inside the Word doc (auto-detected)
    2. ZIP file containing images referenced by IMAGE: filename.png

QUESTION LINE — any of:
    1. text        1) text       Q1. text      Q1) text
    Q 1. text      Question 1.   (1) text      1- text
    1: text        40. text      Q1 text (no delimiter)

OPTION LINE — any of:
    A. text   A) text   A: text   A- text
    a. text   (A) text  A.text (no space)

ANSWER LINE — any of:
    ANSWER: A    Answer A     Ans: B      ANS B
    Answer:C     ANSWER:A     answer = A  ANSWER- A
    Correct Answer: D         The answer is A   Key: B

MARKS LINE — any of:
    MARKS: 1    MARKS: 0.5    MARKS: 1/2    MARKS: 2/4
    POINTS: 2   MARK: .5      score: 1.5    marks 2

INLINE (entire question on one paragraph):
    40. According to ionic theory, acids produce A. OH- ions B. H+ ions C. Na+ ions D. Cl- ions ANSWER: B MARK: 0.5
    Q1. Capital city?A.LagosB.AbujaC.KanoD.IbadanANSWER:B

IMAGE REFERENCE FORMATS (all accepted):
    Standalone line before/after question:
        IMAGE: t1.png
    Inline within question text (stripped automatically):
        Q3. Calculate the angle IMAGE: t1.png A. 48.59 B. 40.55 ANSWER: A
    Image pasted/embedded directly in Word paragraph (auto-detected)
"""

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from io import BytesIO
import re
import hashlib
import zipfile
import os


# ══════════════════════════════════════════════════════════════
#  COMPILED PATTERNS
# ══════════════════════════════════════════════════════════════

RE_QUESTION = re.compile(
    r'^(?:\(?\s*Q(?:uestion)?\s*\.?\s*)?'
    r'\(?\s*(\d+)\s*[.):\-]?\s*\)?\s*'   # already handles no-space
    r'(.+)',
    re.IGNORECASE | re.DOTALL,
)

RE_OPTION_LINE = re.compile(
    r'^\(?([A-Ea-e])[.):\-]\)?\s*(.*)',
    re.IGNORECASE | re.DOTALL,
)

RE_ANSWER = re.compile(
    r'^(?:the\s+)?'
    r'(?:ANSWER|ANS(?:WER)?|CORRECT\s*ANSWER|KEY)'
    r'(?:\s+is)?'
    r'[\s:=\-]*([A-Ea-e])\b'
    r'[\s\(\w\)]*$',   # ✅ allow trailing "(Friendly)" or similar
    re.IGNORECASE,
)

RE_ANSWER_SEARCH = re.compile(
    r'(?:ANSWER|ANS(?:WER)?|CORRECT\s*ANSWER|KEY)'
    r'(?:\s+is)?'
    r'[\s:=\-]*([A-Ea-e])\b',
    re.IGNORECASE,
)

RE_MARKS = re.compile(
    r'^(?:MARKS?|POINTS?|SCORE|MARK|WEIGHT)[:\s=]+(.+)',
    re.IGNORECASE,
)

RE_MARKS_SEARCH = re.compile(
    r'(?:MARKS?|POINTS?|SCORE|MARK|WEIGHT)[:\s=]+([0-9/.]+)',
    re.IGNORECASE,
)

RE_OPT_SPLIT = re.compile(r'(?=[A-Ea-e][.)])')

RE_INLINE_OPT_SPACED = re.compile(r'(?:^|\s)([A-Ea-e])[.)]\s*\S', re.IGNORECASE)
RE_INLINE_OPT_GLUED  = re.compile(r'(?<=[a-z\d?!,;])([A-Ea-e])[.)]', re.IGNORECASE)

# Matches IMAGE: reference anywhere in a line
RE_IMAGE_INLINE = re.compile(r'\s*IMAGE\s*:\s*(\S+)', re.IGNORECASE)


# ══════════════════════════════════════════════════════════════
#  HELPERS
# ══════════════════════════════════════════════════════════════

def _parse_marks(raw):
    if not raw:
        return 1.0
    raw = raw.strip()
    if '/' in raw:
        parts = raw.split('/', 1)
        try:
            return round(float(parts[0].strip()) / float(parts[1].strip()), 4)
        except (ValueError, ZeroDivisionError):
            pass
    try:
        return float(raw)
    except ValueError:
        return 1.0


def _find_inline_options_start(text):
    candidates = []
    for m in RE_INLINE_OPT_SPACED.finditer(text):
        candidates.append((m.start(), m.group(1).upper()))
    for m in RE_INLINE_OPT_GLUED.finditer(text):
        candidates.append((m.start(), m.group(1).upper()))
    if not candidates:
        return -1
    distinct_letters = {c[1] for c in candidates}
    if len(distinct_letters) < 2:
        return -1
    return min(c[0] for c in candidates)


def _is_inline_question(text):
    if not RE_QUESTION.match(text):
        return False
    q_m = RE_QUESTION.match(text)
    rest = text[q_m.start(2):]
    return _find_inline_options_start(rest) >= 0


def _extract_inline_options(options_text):
    options = {}
    answer  = None
    marks   = None

    for chunk in RE_OPT_SPLIT.split(options_text):
        chunk = chunk.strip()
        if not chunk:
            continue

        m = re.match(r'^([A-Ea-e])[.)]\s*(.*)', chunk, re.IGNORECASE | re.DOTALL)
        if not m:
            a = RE_ANSWER_SEARCH.search(chunk)
            if a:
                answer = a.group(1).upper()
            mk = RE_MARKS_SEARCH.search(chunk)
            if mk:
                marks = _parse_marks(mk.group(1))
            continue

        letter      = m.group(1).upper()
        option_text = m.group(2).strip()

        mk = RE_MARKS_SEARCH.search(option_text)
        if mk:
            marks = _parse_marks(mk.group(1))

        a = RE_ANSWER_SEARCH.search(option_text)
        if a:
            answer = a.group(1).upper()
            cut = a.start()
            if mk and mk.start() < cut:
                cut = mk.start()
            option_text = option_text[:cut].strip()
        elif mk:
            option_text = option_text[:mk.start()].strip()

        if option_text:
            options[letter] = option_text

    return options, answer, marks


def _extract_paragraph_image(paragraph, document):
    """
    Extract embedded/pasted image bytes from a Word paragraph.
    Returns (filename, bytes) or (None, None).
    """
    try:
        drawing_els = paragraph._element.findall('.//' + qn('w:drawing'))
        if not drawing_els:
            # Also check for inline images via w:pict
            pict_els = paragraph._element.findall('.//' + qn('w:pict'))
            if not pict_els:
                return None, None

        blip = None
        for drawing in drawing_els:
            blip = drawing.find('.//' + qn('a:blip'))
            if blip is not None:
                break

        # Fallback: check v:imagedata in w:pict
        if blip is None:
            for pict in paragraph._element.findall('.//' + qn('w:pict')):
                imagedata = pict.find('.//' + qn('v:imagedata'))
                if imagedata is not None:
                    r_id = imagedata.get(qn('r:id'))
                    if r_id:
                        image_part = paragraph.part.related_parts.get(r_id)
                        if image_part:
                            image_bytes = image_part.blob
                            ext = image_part.content_type.split('/')[-1]
                            if ext == 'jpeg':
                                ext = 'jpg'
                            fname = f"embed_{hashlib.md5(image_bytes[:64]).hexdigest()[:10]}.{ext}"
                            return fname, image_bytes
            return None, None

        r_id = blip.get(qn('r:embed'))
        if not r_id:
            return None, None

        image_part = paragraph.part.related_parts.get(r_id)
        if image_part is None:
            return None, None

        image_bytes = image_part.blob
        ext = image_part.content_type.split('/')[-1]
        if ext == 'jpeg':
            ext = 'jpg'
        fname = f"embed_{hashlib.md5(image_bytes[:64]).hexdigest()[:10]}.{ext}"
        return fname, image_bytes

    except Exception:
        return None, None


def _load_zip_images(zip_path):
    """
    Load all images from a ZIP file into a dict keyed by bare filename.
    e.g. "folder/t1.png" is stored under both "t1.png" and "folder/t1.png".
    Lookup is also case-insensitive.
    Returns dict: {filename_lower: (original_filename, bytes)}
    """
    images = {}
    try:
        with zipfile.ZipFile(zip_path, 'r') as zf:
            for name in zf.namelist():
                if name.endswith('/'):
                    continue
                # Only load image files
                ext = os.path.splitext(name)[1].lower()
                if ext not in ('.png', '.jpg', '.jpeg', '.gif', '.bmp', '.webp', '.svg'):
                    continue
                data = zf.read(name)
                bare = os.path.basename(name)
                # Store by bare name (case-insensitive)
                images[bare.lower()] = (bare, data)
                # Also store by full path (case-insensitive)
                images[name.lower()] = (bare, data)
    except Exception as e:
        pass
    return images


def _lookup_zip_image(zip_images, ref):
    """
    Look up an image reference in the ZIP dict.
    Returns (filename, bytes) or (None, None).
    Tries exact match first, then case-insensitive.
    """
    if not zip_images or not ref:
        return None, None
    # Try exact (lowercased)
    result = zip_images.get(ref.lower())
    if result:
        return result
    # Try bare filename only
    bare = os.path.basename(ref).lower()
    result = zip_images.get(bare)
    if result:
        return result
    return None, None


# ══════════════════════════════════════════════════════════════
#  MAIN PARSER CLASS
# ══════════════════════════════════════════════════════════════

class WordQuestionParser:
    """
    Robust parser for Word documents containing CBT questions.

    Supports:
      - Images embedded/pasted directly in the Word document
      - Images in a ZIP file referenced by IMAGE: filename.png
      - Both simultaneously (embedded takes priority)

    Usage:
        # Embedded images only
        parser = WordQuestionParser("exam.docx")

        # ZIP images only
        parser = WordQuestionParser("exam.docx", image_zip="images.zip")

        # Both (embedded auto-detected, ZIP as fallback)
        parser = WordQuestionParser("exam.docx", image_zip="images.zip")

        questions = parser.parse()
    """

    def __init__(self, file_path_or_stream, image_zip=None):
        """
        Args:
            file_path_or_stream: Path to .docx file or file-like stream.
            image_zip: Optional path to a ZIP file containing question images.
                       Can also be a file-like stream (e.g. from request.FILES).
        """
        self.document  = Document(file_path_or_stream)
        self.questions = []
        self.current_passage       = None
        self.current_passage_title = None
        self.current_passage_group = None
        self._pending_image_ref    = None
        self._pending_image_cap    = None
        self._pending_emb_fname    = None
        self._pending_emb_bytes    = None

        # Load ZIP images at startup
        self._zip_images = {}
        if image_zip:
            self._zip_images = _load_zip_images(image_zip)

    # ─────────────────────────────────────────────────────────
    #  Public API
    # ─────────────────────────────────────────────────────────

    def parse(self):
        current_question = None
        current_options  = {}

        for paragraph in self.document.paragraphs:
            raw_text = paragraph.text.strip()

            # ── Step 1: Extract embedded image from this paragraph ────────
            # Do NOT attach yet — we need to know which question owns it.
            emb_fname, emb_bytes = _extract_paragraph_image(paragraph, self.document)

            # ── Step 2: Extract any inline IMAGE: reference from the text ──
            # e.g. "Q3. Calculate angle IMAGE: t1.png A. 48 B. 40 ANSWER: A"
            inline_img_ref = None
            inline_match = RE_IMAGE_INLINE.search(raw_text)
            if inline_match:
                inline_img_ref = inline_match.group(1).strip()

            # Strip IMAGE: and CAPTION: text markers from the working text
            text = RE_IMAGE_INLINE.sub('', raw_text).strip()
            text = re.sub(r'\s*CAPTION\s*:\s*.+$', '', text, flags=re.IGNORECASE).strip()

            # ── Step 3: Resolve which image belongs to this paragraph ──────
            # Priority: embedded image > ZIP lookup via inline ref > ZIP lookup via pending ref
            resolved_fname = None
            resolved_bytes = None

            if emb_fname and emb_bytes:
                # Directly embedded in the Word doc
                resolved_fname = emb_fname
                resolved_bytes = emb_bytes
            elif inline_img_ref:
                # TEXT reference like IMAGE: t1.png — look up in ZIP
                zf, zb = _lookup_zip_image(self._zip_images, inline_img_ref)
                if zf and zb:
                    resolved_fname = zf
                    resolved_bytes = zb
                else:
                    # ZIP not provided or file not found — store ref only
                    resolved_fname = inline_img_ref

            # ── Step 4: Handle image-only paragraphs (no text content) ────
            if not text:
                # This paragraph is purely an image or blank line.
                # Attach to the current question if we have one.
                if resolved_fname and current_question:
                    if not current_question.get('image_reference'):
                        current_question['image_reference'] = resolved_fname
                        current_question['image_bytes']     = resolved_bytes
                elif resolved_fname and not current_question:
                    # Image appeared before the question — hold as pending
                    self._pending_emb_fname = resolved_fname
                    self._pending_emb_bytes = resolved_bytes
                continue

            # ── Passage markers ───────────────────────────────────────────
            pg = re.match(r'^PASSAGE[_\s]?GROUP[:\s]+(.+)', text, re.IGNORECASE)
            if pg:
                self.current_passage_group = pg.group(1).strip()
                continue

            pt = re.match(r'^PASSAGE[:\s]+(.+)', text, re.IGNORECASE)
            if pt:
                self.current_passage_title = pt.group(1).strip()
                continue

            if re.match(r'^START[_\s]?PASSAGE', text, re.IGNORECASE):
                self.current_passage = []
                continue

            if re.match(r'^END[_\s]?PASSAGE', text, re.IGNORECASE):
                if isinstance(self.current_passage, list):
                    self.current_passage = '\n'.join(self.current_passage)
                continue

            if isinstance(self.current_passage, list):
                self.current_passage.append(text)
                continue

            # ── Standalone IMAGE: line (after stripping, text is empty) ───
            # Handled above in the "no text" block — but if the line was
            # ONLY "IMAGE: t1.png" with nothing else, text is now empty.
            # The inline_img_ref was captured; resolved above. Nothing to do here.

            # ── Standalone CAPTION: line ──────────────────────────────────
            cap = re.match(r'^CAPTION[:\s]+(.+)', text, re.IGNORECASE)
            if cap:
                caption_val = cap.group(1).strip()
                if current_question:
                    current_question['image_caption'] = caption_val
                else:
                    self._pending_image_cap = caption_val
                continue

            # ── Marks line ────────────────────────────────────────────────
            mk = RE_MARKS.match(text)
            if mk and current_question:
                current_question['marks'] = _parse_marks(mk.group(1))
                continue

            # ── Short-answer marker ───────────────────────────────────────
            if re.match(r'^(?:SHORT\s*ANSWER|SA)[:\s]*$', text, re.IGNORECASE) and current_question:
                current_question['question_type'] = 'SA'
                continue

            sa = re.match(r'^(?:ANSWER|ANS)[:\s]+(.+)', text, re.IGNORECASE)
            if sa and current_question and current_question.get('question_type') == 'SA':
                current_question['correct_answer'] = sa.group(1).strip()
                continue

            # ══ INLINE QUESTION ══════════════════════════════════════════
            if _is_inline_question(text):
                # Save previous question
                if current_question:
                    current_question['options'] = current_options
                    self.questions.append(current_question)
                    current_question = None
                    current_options  = {}

                q_m       = RE_QUESTION.match(text)
                full_rest = text[q_m.start(2):]
                opt_idx   = _find_inline_options_start(full_rest)

                question_text = full_rest[:opt_idx].strip() if opt_idx > 0 else full_rest.strip()
                options_text  = full_rest[opt_idx:]         if opt_idx > 0 else ''

                opts, ans, inline_marks = _extract_inline_options(options_text)

                q = self._new_question(q_m.group(1), question_text)
                q['options'] = opts
                if ans:
                    q['correct_option'] = ans
                if inline_marks is not None:
                    q['marks'] = inline_marks

                # Attach image from THIS paragraph (embedded or ZIP reference)
                self._attach_image(q, resolved_fname, resolved_bytes)

                self.questions.append(q)
                continue

            # ══ NORMAL QUESTION LINE ════════════════════════════════════
            q_m = RE_QUESTION.match(text)
            if q_m:
                # Save previous question
                if current_question:
                    current_question['options'] = current_options
                    self.questions.append(current_question)

                q_html = self._paragraph_to_html(paragraph, strip_q_prefix=True)
                current_question = self._new_question(q_m.group(1), q_html)
                current_options  = {}

                # Attach image from THIS paragraph
                self._attach_image(current_question, resolved_fname, resolved_bytes)
                continue

            # ── Option line ───────────────────────────────────────────────
            opt_m = RE_OPTION_LINE.match(text)
            if opt_m and current_question:
                letter   = opt_m.group(1).upper()
                opt_html = self._paragraph_to_html(paragraph, strip_opt_prefix=True)
                current_options[letter] = opt_html

                # Image on same line as an option — attach to parent question
                if resolved_fname and not current_question.get('image_reference'):
                    self._attach_image(current_question, resolved_fname, resolved_bytes)
                continue

            # ── Answer line ───────────────────────────────────────────────
            ans_m = RE_ANSWER.match(text)
            if ans_m and current_question:
                current_question['correct_option'] = ans_m.group(1).upper()
                if len(current_options) == 2:
                    vals = [str(v).lower() for v in current_options.values()]
                    if 'true' in vals and 'false' in vals:
                        current_question['question_type'] = 'TF'
                continue

        # Save last question
        if current_question:
            current_question['options'] = current_options
            self.questions.append(current_question)

        return self.questions

    # ─────────────────────────────────────────────────────────
    #  Helpers
    # ─────────────────────────────────────────────────────────

    def _attach_image(self, question, fname, data):
        """
        Attach image to a question dict.
        Does nothing if the question already has an image.
        If data is None but fname is set, stores reference only
        (useful when ZIP not provided but filename is known).
        """
        if not fname:
            return
        if question.get('image_reference'):
            return  # already has one — don't overwrite
        question['image_reference'] = fname
        question['image_bytes']     = data  # may be None if ZIP not supplied

    def _new_question(self, number, text):
        """
        Build a fresh question dict, consuming any pending image/caption.
        Pending embedded image (from a paragraph before the question line)
        is also consumed here.
        """
        # Resolve pending image — embedded takes priority over text reference
        pending_ref   = getattr(self, '_pending_image_ref', None)
        pending_cap   = getattr(self, '_pending_image_cap', None)
        pending_efname = getattr(self, '_pending_emb_fname', None)
        pending_ebytes = getattr(self, '_pending_emb_bytes', None)

        # Choose best pending image source
        if pending_efname:
            img_ref   = pending_efname
            img_bytes = pending_ebytes
        elif pending_ref:
            # Try ZIP lookup for the text reference
            zf, zb = _lookup_zip_image(self._zip_images, pending_ref)
            img_ref   = zf if zf else pending_ref
            img_bytes = zb
        else:
            img_ref   = None
            img_bytes = None

        q = {
            'question_number': number,
            'question_text':   text,
            'question_type':   'MCQ',
            'marks':           1.0,
            'correct_option':  None,
            'correct_answer':  None,
            'passage':         self.current_passage if isinstance(self.current_passage, str) else None,
            'passage_title':   self.current_passage_title,
            'passage_group':   self.current_passage_group,
            'image_reference': img_ref,
            'image_caption':   pending_cap,
            'image_bytes':     img_bytes,
        }

        # Clear all pending state
        self._pending_image_ref = None
        self._pending_image_cap = None
        self._pending_emb_fname = None
        self._pending_emb_bytes = None

        return q

    def _paragraph_to_html(self, paragraph, strip_q_prefix=False, strip_opt_prefix=False):
        """Convert paragraph runs to HTML preserving bold/italic/super/subscript."""
        html_parts      = []
        prefix_stripped = False

        for run in paragraph.runs:
            text = run.text
            if not prefix_stripped:
                if strip_q_prefix:
                    text = re.sub(
                        r'^(?:\(?\s*Q(?:uestion)?\s*\.?\s*)?\(?\s*\d+\s*[.):\-]?\s*\)?\s*',
                        '', text, flags=re.IGNORECASE
                    )
                    prefix_stripped = True
                elif strip_opt_prefix:
                    text = re.sub(
                        r'^\(?[A-Ea-e][.):\-]\)?\s*',
                        '', text, flags=re.IGNORECASE
                    )
                    prefix_stripped = True

            if not text:
                continue

            html = text
            if run.font.subscript:   html = f'<sub>{html}</sub>'
            if run.font.superscript: html = f'<sup>{html}</sup>'
            if run.bold:             html = f'<strong>{html}</strong>'
            if run.italic:           html = f'<em>{html}</em>'
            if run.underline:        html = f'<u>{html}</u>'
            html_parts.append(html)

        return ''.join(html_parts)

    def validate_questions(self):
        errors = []
        for idx, q in enumerate(self.questions, 1):
            if not q.get('question_text'):
                errors.append(f"Question {idx}: Missing question text")
            if q['question_type'] in ['MCQ', 'TF']:
                if not q.get('options') or len(q['options']) < 2:
                    errors.append(f"Question {idx}: Needs at least 2 options")
                if not q.get('correct_option'):
                    errors.append(f"Question {idx}: Missing correct answer")
                if q.get('correct_option') and q['correct_option'] not in q.get('options', {}):
                    errors.append(
                        f"Question {idx}: Answer '{q['correct_option']}' "
                        f"not in options {list(q.get('options', {}).keys())}"
                    )
            elif q['question_type'] == 'SA':
                if not q.get('correct_answer'):
                    errors.append(f"Question {idx}: Missing answer")
        return errors

    def get_statistics(self):
        stats = {
            'total': len(self.questions), 'mcq': 0, 'tf': 0, 'sa': 0,
            'total_marks': 0.0, 'with_passage': 0, 'with_image': 0,
            'passage_groups': set(),
        }
        for q in self.questions:
            t = q['question_type']
            if t == 'MCQ':  stats['mcq'] += 1
            elif t == 'TF': stats['tf']  += 1
            elif t == 'SA': stats['sa']  += 1
            stats['total_marks'] += q.get('marks', 1.0)
            if q.get('passage'):         stats['with_passage'] += 1
            if q.get('image_reference'): stats['with_image']   += 1
            if q.get('passage_group'):   stats['passage_groups'].add(q['passage_group'])
        stats['passage_groups'] = len(stats['passage_groups'])
        return stats

    def get_image_as_base64(self, question):
        """
        Returns a data URI string for the question's image, ready for <img src="...">.
        Returns None if the question has no image bytes.

        Usage in template:
            img_src = parser.get_image_as_base64(question)
            # <img src="{{ img_src }}">
        """
        import base64
        data = question.get('image_bytes')
        ref  = question.get('image_reference', '')
        if not data:
            return None
        ext  = os.path.splitext(ref)[1].lower().lstrip('.')
        mime = {
            'png':  'image/png',
            'jpg':  'image/jpeg',
            'jpeg': 'image/jpeg',
            'gif':  'image/gif',
            'bmp':  'image/bmp',
            'webp': 'image/webp',
            'svg':  'image/svg+xml',
        }.get(ext, 'image/png')
        return f"data:{mime};base64,{base64.b64encode(data).decode()}"


# ══════════════════════════════════════════════════════════════
#  TEMPLATE GENERATOR
# ══════════════════════════════════════════════════════════════

def create_word_template():
    doc = Document()

    t = doc.add_paragraph('CBT QUESTIONS TEMPLATE')
    t.runs[0].bold = True
    t.runs[0].font.size = Pt(16)

    doc.add_paragraph('All formats below are accepted:')
    for line in [
        'Questions : 1. text  |  1) text  |  Q1. text  |  (1) text  |  Q1 text',
        'Options   : A. text  |  A) text  |  (A) text  |  a. text   |  A.text',
        'Answers   : ANSWER: A  |  Answer A  |  Ans: B  |  The answer is C',
        'Marks     : MARKS: 1   |  MARKS: 0.5  |  MARKS: 1/2  |  POINTS: 2',
        'Inline    : 40. Question A. Opt1 B. Opt2 C. Opt3 D. Opt4 ANSWER: A MARKS: 0.5',
        'Images    : IMAGE: t1.png  (on its own line, or inline within question text)',
    ]:
        doc.add_paragraph(line)

    doc.add_paragraph()
    doc.add_paragraph('=== Normal format ===').runs[0].bold = True
    doc.add_paragraph('1. The capital of Nigeria is')
    doc.add_paragraph('A. Lagos')
    doc.add_paragraph('B. Abuja')
    doc.add_paragraph('C. Kano')
    doc.add_paragraph('D. Ibadan')
    doc.add_paragraph('ANSWER: B')
    doc.add_paragraph('MARKS: 0.5')

    doc.add_paragraph()
    doc.add_paragraph('=== With image reference (ZIP) ===').runs[0].bold = True
    doc.add_paragraph('IMAGE: figure1.png')
    doc.add_paragraph('2. Calculate the angle marked b in the figure above')
    doc.add_paragraph('A. 48.59')
    doc.add_paragraph('B. 40.55')
    doc.add_paragraph('C. 45.59')
    doc.add_paragraph('D. 26.56')
    doc.add_paragraph('ANSWER: A')

    doc.add_paragraph()
    doc.add_paragraph('=== Inline format with image ===').runs[0].bold = True
    doc.add_paragraph('3. Calculate angle b IMAGE: t1.png A. 48.59 B. 40.55 C. 45.59 D. 26.56 ANSWER: A MARK: 0.5')

    doc.add_paragraph()
    doc.add_paragraph('=== Inline format ===').runs[0].bold = True
    doc.add_paragraph('40. According to ionic theory, acids produce A. OH- ions B. H+ ions C. Na+ ions D. Cl- ions ANSWER: B MARK: 0.5')

    doc.add_paragraph()
    doc.add_paragraph('=== Fraction marks ===').runs[0].bold = True
    doc.add_paragraph('4. What is H2O?')
    doc.add_paragraph('A. Salt')
    doc.add_paragraph('B. Water')
    doc.add_paragraph('C. Sugar')
    doc.add_paragraph('D. Oil')
    doc.add_paragraph('ANSWER: B')
    doc.add_paragraph('MARKS: 1/2')

    return doc


# ══════════════════════════════════════════════════════════════
#  USAGE EXAMPLES
# ══════════════════════════════════════════════════════════════
"""
─────────────────────────────────────────────────────────────
1. EMBEDDED IMAGES ONLY (images pasted directly into Word)
─────────────────────────────────────────────────────────────
parser = WordQuestionParser("exam.docx")
questions = parser.parse()

─────────────────────────────────────────────────────────────
2. ZIP IMAGES ONLY (IMAGE: filename.png references in text)
─────────────────────────────────────────────────────────────
parser = WordQuestionParser("exam.docx", image_zip="images.zip")
questions = parser.parse()

─────────────────────────────────────────────────────────────
3. BOTH (embedded auto-detected, ZIP as fallback)
─────────────────────────────────────────────────────────────
parser = WordQuestionParser("exam.docx", image_zip="images.zip")
questions = parser.parse()

─────────────────────────────────────────────────────────────
4. DJANGO VIEW EXAMPLE
─────────────────────────────────────────────────────────────
def upload_exam(request):
    docx_file = request.FILES['docx']
    zip_file  = request.FILES.get('zip')   # optional

    parser    = WordQuestionParser(docx_file, image_zip=zip_file)
    questions = parser.parse()
    errors    = parser.validate_questions()

    # Attach base64 image to each question for template rendering
    for q in questions:
        q['image_src'] = parser.get_image_as_base64(q)  # None if no image

    return render(request, 'exam.html', {'questions': questions, 'errors': errors})

─────────────────────────────────────────────────────────────
5. DJANGO TEMPLATE
─────────────────────────────────────────────────────────────
{% for q in questions %}
  <div class="question">
    <p>{{ q.question_number }}. {{ q.question_text|safe }}</p>

    {% if q.image_src %}
      <img src="{{ q.image_src }}" alt="Question image" style="max-width:100%;margin:8px 0;">
    {% endif %}

    {% for letter, text in q.options.items %}
      <label>
        <input type="radio" name="q{{ q.question_number }}" value="{{ letter }}">
        {{ letter }}. {{ text|safe }}
      </label>
    {% endfor %}
  </div>
{% endfor %}
"""